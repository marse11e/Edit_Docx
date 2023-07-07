from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text

def replace_text_in_tables(doc, old_text, new_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)
                    if old_text in paragraph.text:
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                text = inline[i].text.replace(old_text, new_text)
                                inline[i].text = text

# Открываем документ .docx
doc = Document('ASU_Силлабус_ШАБЛОН_на_2022_2023_уч_г_АНГ.docx')

# Заменяем значения переменных
dean_name = "Dean of School/Center Name/Surname"
day = "10"
month = "July"
year = "2023"
code_and_course = "123 Course"
ects_and_hourse = "3 ECTS Total hours: 90 Classroom hours: 30 hours  Independent work (IWST, IWS): 60 hours"
level_of_course = "Bachelor degree "


replace_text(doc, "{ dean_of_school_center }", dean_name)
replace_text(doc, "{ day }", day)
replace_text(doc, "{ month }", month)
replace_text(doc, "{ year }", year)
replace_text(doc, "{ code_and_course }", code_and_course)
replace_text(doc, "{ level_of_course }", level_of_course)

replace_text_in_tables(doc, "{ dean_of_school_center }", dean_name)
replace_text_in_tables(doc, "{ day }", day)
replace_text_in_tables(doc, "{ month }", month)
replace_text_in_tables(doc, "{ year }", year)
replace_text_in_tables(doc, "{ code_and_course }", code_and_course)
replace_text_in_tables(doc, "{ ects_and_hourse }", ects_and_hourse)
replace_text_in_tables(doc, "{ level_of_course }", level_of_course)

# Сохраняем изменения
doc.save('q.docx')
